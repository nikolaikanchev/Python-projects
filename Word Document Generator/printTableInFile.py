from docx.shared import Pt
from docx.shared import Cm, Inches

def printTableInFile(row,doc):
    records=[['МОДУЛ', row[1]],
            ['КАТЕГОРИЯ', "", 'НИВО НА ТРУДНОСТ', row[5]],
            ['ПОДМОДУЛ', row[2], 'ТЕМА', row[3]],
            ['Въпрос', row[6]],
            ['Верен отговор', row[7]],
            ['Грешен отговор', row[8]],
            ['Грешен отговор', row[9]]]

    table = doc.add_table(7,4)
    table.style='Table Grid'

    for cell in table.columns[0].cells:
        cell.width = Cm(4)
    for cell in table.columns[1].cells:
        cell.width = Cm(3.8)
    for cell in table.columns[2].cells:
        cell.width = Cm(6)
    for cell in table.columns[3].cells:
        cell.width = Cm(1.7)

    row1Cells = table.rows[0].cells
    row1Cells[0].merge(row1Cells[1]).text = records[0][0]
    row1Cells[0].paragraphs[0].runs[0].font.bold=True
    row1Cells[0].paragraphs[0].paragraph_format.space_before = Pt(6)
    row1Cells[0].paragraphs[0].paragraph_format.space_after = Pt(6)
    row1Cells[2].merge(row1Cells[3]).text = str(records[0][1])
    row1Cells[2].paragraphs[0].runs[0].font.bold=True
    row1Cells[2].paragraphs[0].paragraph_format.space_before = Pt(6)
    row1Cells[2].paragraphs[0].paragraph_format.space_after = Pt(6)

    row2Cells = table.rows[1].cells
    row2Cells[0].text = records[1][0]
    row2Cells[0].paragraphs[0].runs[0].font.bold=True
    row2Cells[1].text = str(records[1][1])
    row2Cells[2].text = records[1][2]
    row2Cells[2].paragraphs[0].runs[0].font.bold=True
    row2Cells[3].text = str(records[1][3])

    row3Cells = table.rows[2].cells
    row3Cells[0].text = records[2][0]
    row3Cells[0].paragraphs[0].runs[0].font.bold=True
    row3Cells[1].text = str(records[2][1])
    row3Cells[2].text = records[2][2]
    row3Cells[2].paragraphs[0].runs[0].font.bold=True
    row3Cells[3].text = str(records[2][3])

    row4Cells = table.rows[3].cells
    row4Cells[0].text = records[3][0]
    row4Cells[1].merge(row4Cells[2]).merge(row4Cells[3]).text = records[3][1]

    row5Cells = table.rows[4].cells
    row5Cells[0].text = records[4][0]
    row5Cells[1].merge(row5Cells[2]).merge(row5Cells[3]).text = records[4][1]

    row6Cells = table.rows[5].cells
    row6Cells[0].text = records[5][0]
    row6Cells[1].merge(row6Cells[2]).merge(row6Cells[3]).text = records[5][1]

    row7Cells = table.rows[6].cells
    row7Cells[0].text = records[6][0]
    row7Cells[1].merge(row7Cells[2]).merge(row7Cells[3]).text = records[6][1]

    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.paragraph_format.space_before = Pt(3)
                paragraph.paragraph_format.space_after = Pt(6)
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(14)
                    font.name = 'Times New Roman'

    paragraph = doc.add_paragraph('')